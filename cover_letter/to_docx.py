from docx import Document, shared

from constants.cover_letter_const import ATTRIBUTES, BODY1, BODY2, CONCLUSION, INTRO
from cover_letter.utils import format_date, skip_line


# MARK: INITIALIZE .DOCX
def init_docx():
    docx = Document()
    docx.styles["Normal"].font.name = "Times New Roman"
    docx.styles["Normal"].font.size = shared.Pt(12)
    docx.styles["Normal"].paragraph_format.line_spacing = 1
    docx.styles["Normal"].paragraph_format.space_before = shared.Pt(0)
    docx.styles["Normal"].paragraph_format.space_after = shared.Pt(0)
    docx.sections[0].left_margin = shared.Inches(1)
    docx.sections[0].right_margin = shared.Inches(1)
    docx.sections[0].top_margin = shared.Inches(1)
    docx.sections[0].bottom_margin = shared.Inches(1)
    docx.sections[0].page_width = shared.Inches(8.5)
    docx.sections[0].page_height = shared.Inches(11)
    return docx


# MARK: WRITE .DOCX
def write_docx(
    docx: Document,
    company: str,
    role: str,
    languages: dict,
    technologies: dict,
    attribute: int,
    hiring_manager: str,
    organization: str | None,
    company_different: str | None,
):
    # info: Personal Information
    docx.add_paragraph("Michael Beebe")
    docx.add_paragraph("New York, NY")
    docx.add_paragraph("(847) 274-3448")
    docx.add_paragraph("michaelbeebe1031@gmail.com")
    skip_line(docx)

    # info: Date
    docx.add_paragraph(format_date())
    skip_line(docx)

    # info: Company Information
    if not company_different:
        company_different = company

    if hiring_manager != "Hiring Team":
        docx.add_paragraph(hiring_manager)
        docx.add_paragraph("Hiring Manager")
        docx.add_paragraph(company_different)
        skip_line(docx)
    elif organization:
        docx.add_paragraph(company_different)
        docx.add_paragraph(organization)
        skip_line(docx)

    # info: Salutation
    docx.add_paragraph(f"Dear {hiring_manager},")
    skip_line(docx)

    # info: Introduction
    docx.add_paragraph(
        INTRO.replace("[COMPANY]", company)
        .replace("[ROLE]", role)
        .replace("[LANGUAGES]", languages["required"])
        .replace("[COMPANY_ATTRIBUTE]", ATTRIBUTES[attribute]["company"])
        .replace("[PERSONAL_ATTRIBUTE]", ATTRIBUTES[attribute]["personal"])
    )
    skip_line(docx)

    # info: Body 1
    docx.add_paragraph(BODY1)
    skip_line(docx)

    # info: Body 2
    docx.add_paragraph(
        BODY2.replace("[LANGUAGES]", languages["combined"])
        .replace("[COMPANY]", company)
        .replace("[TECHNOLOGIES]", technologies["combined"])
    )
    skip_line(docx)

    # info: Conclusion
    docx.add_paragraph(CONCLUSION.replace("[COMPANY]", company).replace("[ROLE]", role))
    skip_line(docx)

    # info: Signature
    docx.add_paragraph(f"Sincerely,")
    skip_line(docx)
    docx.add_paragraph("Michael Beebe")

    return docx
